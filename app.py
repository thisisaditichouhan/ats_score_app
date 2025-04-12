# import os
# import asyncio
# import streamlit as st
# import docx2txt
# import fitz  # PyMuPDF
# import spacy
# import re

# # ====== Fix tokenizer/env issues for clean logs ======
# os.environ["TOKENIZERS_PARALLELISM"] = "false"
# try:
#     asyncio.get_running_loop()
# except RuntimeError:
#     asyncio.set_event_loop(asyncio.new_event_loop())

# # ====== Load spaCy model (already installed via requirements.txt) ======
# try:
#     nlp = spacy.load("en_core_web_sm")
# except Exception as e:
#     st.error("❌ Failed to load spaCy model 'en_core_web_sm'. Please check installation.")
#     st.stop()

# # ====== Streamlit Page Setup ======
# st.set_page_config(page_title="ATS Resume Score App", layout="wide")
# st.title("📊 ATS Resume Score Calculator with Insights")

# # ====== File Extraction Helpers ======

# def extract_text_from_pdf(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     return "".join([page.get_text() for page in doc])

# def extract_text_from_docx(file):
#     return docx2txt.process(file)

# def extract_text_from_txt(file):
#     return str(file.read(), "utf-8")

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith(".docx"):
#         return extract_text_from_docx(file)
#     elif file.name.endswith(".txt"):
#         return extract_text_from_txt(file)
#     else:
#         return ""

# # ====== Phrase Extraction Using spaCy ======
# def get_phrases(text):
#     doc = nlp(text)
#     keywords = set()
#     for chunk in doc.noun_chunks:
#         phrase = chunk.text.strip().lower()
#         if len(phrase.split()) > 1:
#             keywords.add(phrase)
#     for ent in doc.ents:
#         keywords.add(ent.text.strip().lower())
#     return list(keywords)

# # ====== Similarity Calculation ======
# def calculate_similarity(jd_phrases, resume_text, threshold=0.6):
#     resume_phrases = get_phrases(resume_text)
#     matched = []
#     missing = []

#     for phrase in jd_phrases:
#         phrase_doc = nlp(phrase)
#         similarities = [phrase_doc.similarity(nlp(rp)) for rp in resume_phrases]
#         max_sim = max(similarities) if similarities else 0

#         if max_sim >= threshold:
#             matched.append((phrase, round(max_sim, 2)))
#         else:
#             missing.append((phrase, round(max_sim, 2)))

#     return matched, missing

# # ====== Streamlit UI ======

# st.subheader("📄 Resume Input")
# resume_file = st.file_uploader("Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# resume_text_input = st.text_area("Or paste resume text here:")

# st.subheader("🧾 Job Description Input")
# jd_file = st.file_uploader("Upload JD (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# jd_text_input = st.text_area("Or paste JD text here:")

# if st.button("🔍 Analyze Resume"):
#     with st.spinner("Analyzing resume, please wait..."):

#         resume_text = extract_text(resume_file) if resume_file else resume_text_input
#         jd_text = extract_text(jd_file) if jd_file else jd_text_input

#         if resume_text.strip() and jd_text.strip():
#             jd_phrases = get_phrases(jd_text)
#             matched, missing = calculate_similarity(jd_phrases, resume_text)

#             score = int((len(matched) / len(jd_phrases)) * 100) if jd_phrases else 0

#             st.subheader("📈 ATS Match Score")
#             st.progress(score)
#             st.metric("Match Score", f"{score}/100")

#             st.subheader("✅ Matched Key Phrases")
#             if matched:
#                 st.success(", ".join([kw for kw, _ in matched]))
#             else:
#                 st.warning("No relevant phrases matched the resume.")

#             st.subheader("📌 Suggestions to Improve Resume")
#             if missing:
#                 st.info("Consider including these key phrases (or their concepts):")
#                 for kw, sim in sorted(missing, key=lambda x: x[1], reverse=True):
#                     st.write(f"• {kw}  (similarity: {sim})")
#             else:
#                 st.success("Excellent! Your resume is highly aligned with the job description.")
#         else:
#             st.warning("Please provide both Resume and JD (upload or paste).")

#------------------NEW UPDATED VERSION--------------------------------
# import streamlit as st
# import docx2txt
# import fitz  # PyMuPDF
# import spacy
# import re

# # Load NLP model
# nlp = spacy.load("en_core_web_md")

# # Streamlit Page Setup
# st.set_page_config(page_title="ATS Resume Score App", layout="wide")
# st.title("📊 ATS Resume Score Calculator with Insights")

# # ========== HELPERS ==========

# def extract_text_from_pdf(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text

# def extract_text_from_docx(file):
#     return docx2txt.process(file)

# def extract_text_from_txt(file):
#     return str(file.read(), "utf-8")

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith(".docx"):
#         return extract_text_from_docx(file)
#     elif file.name.endswith(".txt"):
#         return extract_text_from_txt(file)
#     else:
#         return ""

# # New phrase-based extraction using spaCy noun_chunks and entities
# def get_phrases(text):
#     doc = nlp(text)
#     keywords = set()
#     for chunk in doc.noun_chunks:
#         phrase = chunk.text.strip().lower()
#         if len(phrase.split()) > 1:
#             keywords.add(phrase)
#     for ent in doc.ents:
#         keywords.add(ent.text.strip().lower())
#     return list(keywords)

# def calculate_similarity(jd_phrases, resume_text, threshold=0.6):
#     resume_doc = nlp(resume_text)
#     matched = []
#     missing = []
#     for phrase in jd_phrases:
#         phrase_doc = nlp(phrase)
#         sim = phrase_doc.similarity(resume_doc)
#         if sim >= threshold:
#             matched.append((phrase, round(sim, 2)))
#         else:
#             missing.append((phrase, round(sim, 2)))
#     return matched, missing

# # ========== UI ==========

# st.subheader("📄 Resume Input")
# resume_file = st.file_uploader("Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# resume_text_input = st.text_area("Or paste resume text here:")

# st.subheader("🧾 Job Description Input")
# jd_file = st.file_uploader("Upload JD (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# jd_text_input = st.text_area("Or paste JD text here:")

# if st.button("🔍 Analyze Resume"):
#     with st.spinner("Analyzing resume, please wait..."):
#         # Load resume
#         resume_text = ""
#         if resume_file:
#             resume_text = extract_text(resume_file)
#         elif resume_text_input:
#             resume_text = resume_text_input

#         # Load JD
#         jd_text = ""
#         if jd_file:
#             jd_text = extract_text(jd_file)
#         elif jd_text_input:
#             jd_text = jd_text_input

#         if resume_text.strip() and jd_text.strip():
#             jd_phrases = get_phrases(jd_text)
#             matched, missing = calculate_similarity(jd_phrases, resume_text)

#             score = int((len(matched) / len(jd_phrases)) * 100) if jd_phrases else 0

#             # Result Section
#             st.subheader("📈 ATS Match Score")
#             st.progress(score)
#             st.metric("Match Score", f"{score}/100")

#             st.subheader("✅ Matched Key Phrases")
#             if matched:
#                 st.success(", ".join([kw for kw, _ in matched]))
#             else:
#                 st.warning("No relevant phrases matched the resume.")

#             st.subheader("📌 Suggestions to Improve Resume")
#             if missing:
#                 st.info("Consider including these key phrases (or their concepts):")
#                 for kw, sim in sorted(missing, key=lambda x: x[1], reverse=True):
#                     st.write(f"• {kw}  (similarity: {sim})")
#             else:
#                 st.success("Excellent! Your resume is highly aligned with the job description.")
#         else:
#             st.warning("Please provide both Resume and JD (upload or paste).")
#-------------------v2 below ---- added integrity check 1 for white font --------------------#
# import sys
# import os
# import asyncio
# import streamlit as st
# import docx2txt
# import fitz  # PyMuPDF
# import spacy
# import re
# from docx import Document
# from docx.shared import RGBColor

# # ====== Fix tokenizer/env issues for clean logs ======
# os.environ["TOKENIZERS_PARALLELISM"] = "false"
# try:
#     asyncio.get_running_loop()
# except RuntimeError:
#     asyncio.set_event_loop(asyncio.new_event_loop())

# # ====== Load spaCy model (already installed via requirements.txt) ======
# try:
#     nlp = spacy.load("en_core_web_sm")
# except Exception as e:
#     st.error("❌ Failed to load spaCy model 'en_core_web_sm'. Please check installation.")
#     st.stop()

# # ====== Streamlit Page Setup ======
# st.set_page_config(page_title="ATS Resume Score App", layout="wide")
# st.title("📊 ATS Resume Score Calculator with Insights")

# # ====== File Extraction Helpers ======

# def extract_text_from_pdf(file):
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     return "".join([page.get_text() for page in doc])

# def extract_text_from_docx(file):
#     return docx2txt.process(file)

# def extract_text_from_txt(file):
#     return str(file.read(), "utf-8")

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith(".docx"):
#         return extract_text_from_docx(file)
#     elif file.name.endswith(".txt"):
#         return extract_text_from_txt(file)
#     else:
#         return ""
    
# # ===== Integrity check =====

# # white check in word doc
# def contains_white_font_docx(file):
#     document = Document(file)
#     for para in document.paragraphs:
#         for run in para.runs:
#             font_color = run.font.color
#             # Check if the font color is explicitly white (255, 255, 255)
#             if font_color and font_color.rgb == RGBColor(255, 255, 255):
#                 return True
#             # Additionally, check if the font color is None (which means it's default or black color)
#             elif font_color is None:
#                 continue
#     return False

# # ==== White check in .pdf =====
# def is_white(color, tolerance=0.05):
#     # Assuming color is a tuple (R, G, B)
#     return all(abs(c - 1.0) <= tolerance for c in color)

# def contains_white_font_pdf(file):
#     file.seek(0)
#     doc = fitz.open(stream=file.read(), filetype="pdf")
#     for page in doc:
#         blocks = page.get_text("dict")["blocks"]
#         for block in blocks:
#             for line in block.get("lines", []):
#                 for span in line.get("spans", []):
#                     # Get the color as an integer and convert it to RGB
#                     color = span.get("color", 0)
#                     # Extract the RGB values (PyMuPDF stores colors as 24-bit integers)
#                     r = ((color >> 16) & 0xFF) / 255.0
#                     g = ((color >> 8) & 0xFF) / 255.0
#                     b = (color & 0xFF) / 255.0
#                     # Check if the color is close to white (255, 255, 255)
#                     if is_white((r, g, b)):
#                         return True
#     return False

# # ====== Phrase Extraction Using spaCy ======
# def get_phrases(text):
#     doc = nlp(text)
#     keywords = set()
#     for chunk in doc.noun_chunks:
#         phrase = chunk.text.strip().lower()
#         if len(phrase.split()) > 1:
#             keywords.add(phrase)
#     for ent in doc.ents:
#         keywords.add(ent.text.strip().lower())
#     return list(keywords)

# # ====== Similarity Calculation ======
# def calculate_similarity(jd_phrases, resume_text, threshold=0.6):
#     resume_phrases = get_phrases(resume_text)
#     matched = []
#     missing = []

#     for phrase in jd_phrases:
#         phrase_doc = nlp(phrase)
#         similarities = [phrase_doc.similarity(nlp(rp)) for rp in resume_phrases]
#         max_sim = max(similarities) if similarities else 0

#         if max_sim >= threshold:
#             matched.append((phrase, round(max_sim, 2)))
#         else:
#             missing.append((phrase, round(max_sim, 2)))

#     return matched, missing

# # ====== Streamlit UI ======

# st.subheader("📄 Resume Input")
# resume_file = st.file_uploader("Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# resume_text_input = st.text_area("Or paste resume text here:")

# st.subheader("🧾 Job Description Input")
# jd_file = st.file_uploader("Upload JD (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
# jd_text_input = st.text_area("Or paste JD text here:")

# if st.button("🔍 Analyze Resume"):
#     with st.spinner("Analyzing resume, please wait..."):

#         # anti-cheat: check for white font before processing
#         if resume_file:
#             if resume_file.name.endswith(".docx") and contains_white_font_docx(resume_file):
#                 st.error("⚠️ Suspicious formatting detected: white font in Word file. Resume rejected.")
#                 st.stop()
#             elif resume_file.name.endswith(".pdf") and contains_white_font_pdf(resume_file):
#                 st.error("⚠️ Suspicious formatting detected: white font in PDF file. Resume rejected.")
#                 st.stop()
        
#         resume_text = extract_text(resume_file) if resume_file else resume_text_input
#         jd_text = extract_text(jd_file) if jd_file else jd_text_input

#         if resume_text.strip() and jd_text.strip():
#             jd_phrases = get_phrases(jd_text)
#             matched, missing = calculate_similarity(jd_phrases, resume_text)

#             score = int((len(matched) / len(jd_phrases)) * 100) if jd_phrases else 0

#             st.subheader("📈 ATS Match Score")
#             st.progress(score)
#             st.metric("Match Score", f"{score}/100")

#             st.subheader("✅ Matched Key Phrases")
#             if matched:
#                 st.success(", ".join([kw for kw, _ in matched]))
#             else:
#                 st.warning("No relevant phrases matched the resume.")

#             st.subheader("📌 Suggestions to Improve Resume")
#             if missing:
#                 st.info("Consider including these key phrases (or their concepts):")
#                 for kw, sim in sorted(missing, key=lambda x: x[1], reverse=True):
#                     st.write(f"• {kw}  (similarity: {sim})")
#             else:
#                 st.success("Excellent! Your resume is highly aligned with the job description.")
#         else:
#             st.warning("Please provide both Resume and JD (upload or paste).")
#---------------------corrected version -----------------------------#

import os
import asyncio
import streamlit as st
import docx2txt
import fitz  # PyMuPDF
import spacy
from docx import Document
from docx.shared import RGBColor

# ========== ENVIRONMENT FIX ==========
os.environ["TOKENIZERS_PARALLELISM"] = "false"
try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())

# ========== Load spaCy ==========
try:
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    st.error("❌ Failed to load spaCy model 'en_core_web_sm'. Please check installation.")
    st.stop()

# ========== Streamlit Setup ==========
st.set_page_config(page_title="ATS Resume Score App", layout="wide")
st.title("📊 ATS Resume Score Calculator with Insights")

# ========== Text Extraction ==========
def extract_text(file, ext):
    if ext == ".pdf":
        return extract_text_from_pdf(file)
    elif ext == ".docx":
        return extract_text_from_docx(file)
    elif ext == ".txt":
        return str(file.read(), "utf-8")
    else:
        return ""

def extract_text_from_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    return "".join([page.get_text() for page in doc])

def extract_text_from_docx(file):
    return docx2txt.process(file)

# ========== White Font Checkers ==========
def contains_white_font_docx(file_bytes):
    document = Document(file_bytes)
    for para in document.paragraphs:
        for run in para.runs:
            font_color = run.font.color
            if font_color and font_color.rgb == RGBColor(255, 255, 255):
                return True
    return False

def contains_white_font_pdf(file_bytes):
    file_bytes.seek(0)
    doc = fitz.open(stream=file_bytes.read(), filetype="pdf")
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    color = span.get("color", 0)
                    r = ((color >> 16) & 0xFF) / 255.0
                    g = ((color >> 8) & 0xFF) / 255.0
                    b = (color & 0xFF) / 255.0
                    if all(abs(c - 1.0) <= 0.05 for c in (r, g, b)):
                        return True
    return False

# ========== Phrase Extraction ==========
def get_phrases(text):
    doc = nlp(text)
    keywords = set()
    for chunk in doc.noun_chunks:
        phrase = chunk.text.strip().lower()
        if len(phrase.split()) > 1:
            keywords.add(phrase)
    for ent in doc.ents:
        keywords.add(ent.text.strip().lower())
    return list(keywords)

# ========== Similarity Logic ==========
def calculate_similarity(jd_phrases, resume_text, threshold=0.6):
    resume_phrases = get_phrases(resume_text)
    matched, missing = [], []

    for phrase in jd_phrases:
        phrase_doc = nlp(phrase)
        similarities = [phrase_doc.similarity(nlp(rp)) for rp in resume_phrases]
        max_sim = max(similarities) if similarities else 0
        if max_sim >= threshold:
            matched.append((phrase, round(max_sim, 2)))
        else:
            missing.append((phrase, round(max_sim, 2)))
    return matched, missing

# ========== Streamlit UI ==========
st.subheader("📄 Resume Input")
resume_file = st.file_uploader("Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
resume_text_input = st.text_area("Or paste resume text here:")

st.subheader("🧾 Job Description Input")
jd_file = st.file_uploader("Upload JD (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
jd_text_input = st.text_area("Or paste JD text here:")

if st.button("🔍 Analyze Resume"):
    with st.spinner("Analyzing resume, please wait..."):

        # Resume logic
        resume_text = ""
        if resume_file:
            ext = os.path.splitext(resume_file.name)[1].lower()
            resume_file.seek(0)
            file_bytes = resume_file.read()
            resume_file.seek(0)

            if ext == ".docx" and contains_white_font_docx(resume_file):
                st.error("⚠️ White font detected in DOCX. Resume rejected.")
                st.stop()
            elif ext == ".pdf" and contains_white_font_pdf(resume_file):
                st.error("⚠️ White font detected in PDF. Resume rejected.")
                st.stop()

            resume_file.seek(0)
            resume_text = extract_text(resume_file, ext)
        elif resume_text_input:
            resume_text = resume_text_input

        # JD logic
        jd_text = ""
        if jd_file:
            ext = os.path.splitext(jd_file.name)[1].lower()
            jd_text = extract_text(jd_file, ext)
        elif jd_text_input:
            jd_text = jd_text_input

        # Validate inputs
        if resume_text.strip() and jd_text.strip():
            jd_phrases = get_phrases(jd_text)
            matched, missing = calculate_similarity(jd_phrases, resume_text)

            score = int((len(matched) / len(jd_phrases)) * 100) if jd_phrases else 0

            st.subheader("📈 ATS Match Score")
            st.progress(score)
            st.metric("Match Score", f"{score}/100")

            st.subheader("✅ Matched Key Phrases")
            if matched:
                st.success(", ".join([kw for kw, _ in matched]))
            else:
                st.warning("No relevant phrases matched the resume.")

            st.subheader("📌 Suggestions to Improve Resume")
            if missing:
                st.info("Consider including these key phrases (or their concepts):")
                for kw, sim in sorted(missing, key=lambda x: x[1], reverse=True):
                    st.write(f"• {kw}  (similarity: {sim})")
            else:
                st.success("Excellent! Your resume is highly aligned with the job description.")
        else:
            st.warning("Please provide both Resume and JD (upload or paste).")

